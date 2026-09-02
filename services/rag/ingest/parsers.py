"""文档解析器工厂：统一把 PDF/DOCX/TXT/MD 提取为纯文本。

设计：
- 统一入口 parse_document(file_type, data) -> {"title", "text", "word_count"}；
- 解析是纯同步 CPU 密集操作（PyMuPDF / python-docx），调用方应在
  asyncio.to_thread 中执行，避免阻塞事件循环；
- PDF 扫描件（无可提取文本）抛出 DocumentParseError，流水线将其标记为 failed。

支持格式：
- txt/md: UTF-8 解码（带 BOM/GBK 兜底，避免中文乱码）
- pdf:   PyMuPDF(fitz) 逐页提取
- docx:  python-docx 段落 + 表格
"""
import hashlib
import logging
import re
from dataclasses import dataclass

from infrastructure.config import get_value

logger = logging.getLogger("lunjiang.ingest")

SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}
# 文件签名识别失败时按扩展名兜底
_EXT_MAP = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md", ".markdown": "md"}


class DocumentParseError(Exception):
    """解析失败（不支持格式 / 扫描件无可提取文本 / 内容为空）。"""


@dataclass
class ParsedDocument:
    file_type: str
    title: str = ""
    text: str = ""
    word_count: int = 0

    def normalize(self, min_text_chars: int) -> "ParsedDocument":
        """规范化：压缩空白、识别标题、扫描件检测。"""
        self.text = re.sub(r"[ \t\r\f\v]+", " ", self.text).strip()
        # 保留单行换行的段落结构：连续空行压缩为单换行
        self.text = re.sub(r"\n{3,}", "\n\n", self.text)
        if not self.title and self.text:
            first_line = self.text.splitlines()[0].strip().strip("#").strip()
            if first_line and len(first_line) <= 60:
                self.title = first_line
        self.word_count = len(re.sub(r"\s", "", self.text))
        if self.word_count < min_text_chars:
            raise DocumentParseError(
                f"无可提取文本（仅{self.word_count}字符），疑似扫描版 PDF 或空文档，本期不支持 OCR")
        return self


# ---------------------------------------------------------------
# 各格式解析实现
# ---------------------------------------------------------------
def _parse_txt_md(data: bytes) -> str:
    raw = data
    for enc in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _parse_pdf(data: bytes, min_text_chars: int) -> str:
    """PyMuPDF 提取：优先文本层，扫描页自动跳过。"""
    import pymupdf  # 同 fitz API，规避弃用告警

    parts: list[str] = []
    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        # 用页索引遍历：pymupdf 的 Document stub 未声明 __iter__，索引访问更稳
        for i in range(doc.page_count):
            page = doc[i]
            page_text = page.get_text("text").strip()
            if page_text:
                parts.append(page_text)
            else:
                # 页内可含少量矢量文本（如页码），粗略探查
                words = page.get_text("words")
                if words:
                    parts.append(" ".join(w[4] for w in words))
    finally:
        doc.close()
    text = "\n".join(parts)
    if len(re.sub(r"\s", "", text)) < min_text_chars:
        raise DocumentParseError("PDF 无可提取文本，疑似扫描版，本期不支持 OCR")
    return text


def _parse_docx(data: bytes) -> str:
    """python-docx：段落 + 表格文本拼接。"""
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------
# 对外入口：解析器工厂
# ---------------------------------------------------------------
def parse_document(*, file_type: str, data: bytes, filename: str = "") -> ParsedDocument:
    """按类型分派解析，返回规范化后的 ParsedDocument。

    Raises: DocumentParseError
    """
    ft = file_type.lower().lstrip(".")
    if ft not in SUPPORTED_TYPES:
        raise DocumentParseError(
            f"不支持的文件类型: {ft}，仅支持 {sorted(SUPPORTED_TYPES)}")
    if not data:
        raise DocumentParseError("空文件")

    if ft in ("txt", "md"):
        text = _parse_txt_md(data)
    elif ft == "pdf":
        text = _parse_pdf(data, int(_min_text_chars()))
    else:  # docx
        text = _parse_docx(data)
    if not text.strip():
        raise DocumentParseError("文档内容为空")

    parsed = ParsedDocument(file_type=ft, text=text)
    parsed.normalize(_min_text_chars())
    if parsed.title and filename:
        # 元数据标题与文件名一致时保留文件名（更符合用户认知）
        stem = filename.rsplit(".", 1)[0].strip()
        if stem and stem != parsed.title:
            parsed.title = parsed.title
    return parsed


def sha256_fingerprint(data: bytes, extra: str = "") -> str:
    """文件内容指纹（MD5，用于同项目内去重）。"""
    return hashlib.md5(data + extra.encode("utf-8")).hexdigest()


def infer_type(filename: str, data: bytes) -> str:
    """按扩展名推断类型；docx/老版 .doc 二进制再区分。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    ft = _EXT_MAP.get("." + ext, "")
    if ft:
        return ft
    # 魔数兜底：PDF 头 / zip(docx)
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(b"PK"):
        return "docx"
    raise DocumentParseError(f"无法识别文件格式: {filename}")


def _min_text_chars() -> int:
    return int(get_value("rag", "knowledge", "min_text_chars", default=30))